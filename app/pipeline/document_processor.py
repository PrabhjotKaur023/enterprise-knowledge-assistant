"""
Document processing pipeline.

This is the messiest part of any RAG system — real documents are dirty.
I spent more time here than on the actual RAG logic. PDFs with scanned
images, DOCX with weird formatting, TXT with mixed encodings — all real
problems I hit while testing.
"""

import logging
import re
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Tuple

from app.core.config import settings

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """A single chunk ready for embedding."""
    chunk_id: str
    document_id: str
    content: str
    metadata: dict = field(default_factory=dict)
    # char positions in original doc — useful for citation highlighting
    char_start: int = 0
    char_end: int = 0


@dataclass
class ProcessedDocument:
    """Result of running a file through the pipeline."""
    document_id: str
    filename: str
    file_type: str
    raw_text: str
    chunks: List[DocumentChunk]
    metadata: dict
    word_count: int
    char_count: int


class DocumentValidator:
    """
    Pre-processing checks before we do expensive work.
    Fail fast is the principle here.
    """

    MAX_SIZE = settings.MAX_FILE_SIZE_MB * 1024 * 1024

    def validate(self, file_path: Path, original_name: str) -> Tuple[bool, str]:
        """Returns (is_valid, error_message)."""
        if not file_path.exists():
            return False, f"File not found: {file_path}"

        size = file_path.stat().st_size
        if size == 0:
            return False, "File is empty."
        if size > self.MAX_SIZE:
            return False, f"File exceeds {settings.MAX_FILE_SIZE_MB}MB limit ({size // 1024 // 1024}MB)."

        ext = original_name.rsplit(".", 1)[-1].lower()
        if ext not in settings.ALLOWED_EXTENSIONS:
            return False, f"Unsupported file type: .{ext}. Allowed: {settings.ALLOWED_EXTENSIONS}"

        return True, ""


class TextExtractor:
    """Extract raw text from supported file types."""

    def extract(self, file_path: Path, file_type: str) -> str:
        extractors = {
            "pdf": self._extract_pdf,
            "docx": self._extract_docx,
            "txt": self._extract_txt,
        }
        extractor = extractors.get(file_type)
        if not extractor:
            raise ValueError(f"No extractor for file type: {file_type}")

        raw = extractor(file_path)
        return self._clean_text(raw)

    def _extract_pdf(self, path: Path) -> str:
        try:
            import PyPDF2
            text_parts = []
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"PDF extraction failed for {path}: {e}")
            raise

    def _extract_docx(self, path: Path) -> str:
        try:
            from docx import Document
            doc = Document(str(path))
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            # Also pull text from tables — easy to miss otherwise
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.error(f"DOCX extraction failed for {path}: {e}")
            raise

    def _extract_txt(self, path: Path) -> str:
        # Try UTF-8 first, fall back to latin-1 for older files
        for encoding in ("utf-8", "latin-1", "cp1252"):
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Could not decode {path} with any supported encoding.")

    def _clean_text(self, text: str) -> str:
        """Basic cleanup that helps chunking quality a lot."""
        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        # Collapse 3+ blank lines to 2
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Strip trailing whitespace on each line
        text = "\n".join(line.rstrip() for line in text.split("\n"))
        # Remove null bytes (sometimes in PDFs)
        text = text.replace("\x00", "")
        return text.strip()


class TextChunker:
    """
    Splits text into overlapping chunks for embedding.

    Tried several strategies: fixed-size, sentence-based, paragraph-based.
    Fixed-size with overlap is the most predictable and works well enough
    for general documents. Sentence-based was better for narrative text
    but slower and harder to tune.
    """

    def __init__(
        self,
        chunk_size: int = settings.CHUNK_SIZE,
        overlap: int = settings.CHUNK_OVERLAP,
    ):
        self.chunk_size = chunk_size
        self.overlap = overlap

    def chunk(self, text: str, document_id: str, metadata: dict) -> List[DocumentChunk]:
        """Split text into overlapping chunks."""
        # Try to split on paragraph boundaries first
        paragraphs = text.split("\n\n")
        chunks = []
        current = ""
        current_start = 0
        global_offset = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                global_offset += 2  # account for \n\n
                continue

            if len(current) + len(para) + 2 <= self.chunk_size:
                current = (current + "\n\n" + para).strip() if current else para
            else:
                if current:
                    chunk = self._make_chunk(
                        current, document_id, len(chunks), metadata,
                        current_start, current_start + len(current)
                    )
                    chunks.append(chunk)
                    # Overlap: keep last N chars of previous chunk
                    overlap_text = current[-self.overlap:] if len(current) > self.overlap else current
                    current = (overlap_text + "\n\n" + para).strip()
                    current_start = global_offset - len(overlap_text)
                else:
                    # Single paragraph > chunk_size — split by words
                    word_chunks = self._split_long_text(para, document_id, len(chunks), metadata, global_offset)
                    chunks.extend(word_chunks)
                    current = ""
                    current_start = global_offset + len(para)

            global_offset += len(para) + 2

        if current.strip():
            chunks.append(
                self._make_chunk(current, document_id, len(chunks), metadata, current_start, current_start + len(current))
            )

        logger.debug(f"Chunked document {document_id} into {len(chunks)} chunks.")
        return chunks

    def _make_chunk(
        self, content: str, doc_id: str, idx: int, metadata: dict, start: int, end: int
    ) -> DocumentChunk:
        chunk_meta = {**metadata, "chunk_index": idx, "chunk_size": len(content)}
        return DocumentChunk(
            chunk_id=f"{doc_id}_chunk_{idx}",
            document_id=doc_id,
            content=content.strip(),
            metadata=chunk_meta,
            char_start=start,
            char_end=end,
        )

    def _split_long_text(
        self, text: str, doc_id: str, start_idx: int, metadata: dict, offset: int
    ) -> List[DocumentChunk]:
        """Word-boundary split for paragraphs that exceed chunk_size."""
        words = text.split()
        chunks = []
        current_words = []
        current_len = 0
        char_pos = offset

        for word in words:
            if current_len + len(word) + 1 > self.chunk_size and current_words:
                content = " ".join(current_words)
                chunks.append(
                    self._make_chunk(content, doc_id, start_idx + len(chunks), metadata, char_pos, char_pos + len(content))
                )
                # Overlap by keeping some words
                overlap_words = current_words[-5:]
                char_pos += len(content) - len(" ".join(overlap_words))
                current_words = overlap_words
                current_len = sum(len(w) + 1 for w in overlap_words)

            current_words.append(word)
            current_len += len(word) + 1

        if current_words:
            content = " ".join(current_words)
            chunks.append(
                self._make_chunk(content, doc_id, start_idx + len(chunks), metadata, char_pos, char_pos + len(content))
            )

        return chunks


class MetadataExtractor:
    """Pull useful metadata from documents for filtering and citation."""

    def extract(self, text: str, filename: str, file_type: str) -> dict:
        word_count = len(text.split())
        # Rough page estimate — not exact but good enough for display
        estimated_pages = max(1, word_count // 250)

        return {
            "filename": filename,
            "file_type": file_type,
            "word_count": word_count,
            "estimated_pages": estimated_pages,
            "char_count": len(text),
            # TODO: extract title from first line / PDF metadata
            # TODO: extract author if present
        }


class DocumentPipeline:
    """
    Orchestrates the full processing flow:
    validate → extract → clean → chunk → metadata.
    """

    def __init__(self):
        self.validator = DocumentValidator()
        self.extractor = TextExtractor()
        self.chunker = TextChunker()
        self.metadata_extractor = MetadataExtractor()

    def process(self, file_path: Path, original_name: str, document_id: Optional[str] = None) -> ProcessedDocument:
        """Run a file through the complete pipeline."""
        doc_id = document_id or str(uuid.uuid4())
        file_type = original_name.rsplit(".", 1)[-1].lower()

        logger.info(f"Processing document: {original_name} (id={doc_id})")

        # 1. Validate
        valid, error = self.validator.validate(file_path, original_name)
        if not valid:
            raise ValueError(f"Validation failed: {error}")

        # 2. Extract text
        raw_text = self.extractor.extract(file_path, file_type)
        if not raw_text.strip():
            raise ValueError("Document appears to be empty or unreadable (possibly a scanned image PDF).")

        # 3. Extract metadata
        metadata = self.metadata_extractor.extract(raw_text, original_name, file_type)
        metadata["document_id"] = doc_id

        # 4. Chunk
        chunks = self.chunker.chunk(raw_text, doc_id, metadata)
        if not chunks:
            raise ValueError("Could not generate any chunks from document content.")

        logger.info(f"Document processed: {len(chunks)} chunks, ~{metadata['word_count']} words")

        return ProcessedDocument(
            document_id=doc_id,
            filename=original_name,
            file_type=file_type,
            raw_text=raw_text,
            chunks=chunks,
            metadata=metadata,
            word_count=metadata["word_count"],
            char_count=metadata["char_count"],
        )
